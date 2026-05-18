"""
Export Utilities - Day 7

Generate professional reports from analysis results:
- DOCX reports (executive summary, detailed analysis)
- HTML reports (web viewing)
- Markdown reports (documentation)
- Summary dashboards
"""

import os
import json
import logging
from typing import Dict, List, Optional
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ReportExporter:
    """
    Export analysis results to various formats.
    
    Supports:
    - DOCX (Word documents)
    - HTML (web viewing)
    - Markdown (documentation)
    """
    
    def __init__(self):
        pass
    
    def export_executive_summary_docx(
        self, 
        metadata: Dict, 
        requirements: List[Dict],
        entities: Dict,
        output_path: str
    ) -> str:
        """
        Generate executive summary DOCX report.
        
        1-2 page summary for leadership with:
        - Document overview
        - Key dates
        - Critical requirements
        - Impact summary
        """
        logger.info("Generating executive summary DOCX...")
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading('Regulatory Analysis', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_heading(metadata.get('title', 'Unknown Regulation'), 1)
        
        # Document Info Box
        doc.add_heading('Document Information', 2)
        
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_data = [
            ('Agency', metadata.get('agency', 'Unknown')),
            ('Type', metadata.get('document_type', 'Unknown')),
            ('Regulation Number', metadata.get('regulation_number', 'Unknown')),
            ('Publication Date', metadata.get('publication_date', 'Unknown')),
            ('Effective Date', metadata.get('effective_date', 'TBD')),
            ('Processed', datetime.now().strftime('%Y-%m-%d %H:%M'))
        ]
        
        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
        
        doc.add_paragraph()  # Space
        
        # Executive Summary
        doc.add_heading('Executive Summary', 2)
        summary_text = metadata.get('summary', 'No summary available')
        doc.add_paragraph(summary_text)
        
        doc.add_paragraph()
        
        # Key Dates
        doc.add_heading('Key Compliance Dates', 2)
        
        all_deadlines = []
        for section in requirements:
            all_deadlines.extend(section.get('deadlines', []))
        
        if all_deadlines:
            # Sort by date
            all_deadlines = sorted(
                [d for d in all_deadlines if d.get('date')],
                key=lambda x: x['date']
            )
            
            for deadline in all_deadlines[:5]:  # Top 5
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{deadline['date']}: ").bold = True
                p.add_run(deadline['action'])
        else:
            doc.add_paragraph("No specific deadlines identified.")
        
        doc.add_paragraph()
        
        # Critical Requirements
        doc.add_heading('Critical Requirements', 2)
        
        all_obligations = []
        for section in requirements:
            all_obligations.extend(section.get('obligations', []))
        
        # Filter for high/critical severity
        critical = [o for o in all_obligations if o.get('severity') in ['high', 'critical']]
        
        if critical:
            for i, obligation in enumerate(critical[:5], 1):
                doc.add_heading(f"Requirement {i}", 3)
                doc.add_paragraph(obligation.get('plain_language', obligation.get('text', 'N/A')))
                
                details = doc.add_paragraph()
                details.add_run('Who must comply: ').bold = True
                details.add_run(obligation.get('who_must_comply', 'Not specified'))
                
                if obligation.get('deadline'):
                    deadline_p = doc.add_paragraph()
                    deadline_p.add_run('Deadline: ').bold = True
                    deadline_p.add_run(obligation.get('deadline', 'None'))
        else:
            doc.add_paragraph("See detailed requirements in full analysis.")
        
        # Save
        doc.save(output_path)
        logger.info(f"✅ Saved executive summary: {output_path}")
        
        return output_path
    
    def export_detailed_analysis_docx(
        self,
        metadata: Dict,
        requirements: List[Dict],
        entities: Dict,
        structure: Dict,
        output_path: str
    ) -> str:
        """
        Generate detailed analysis DOCX report.
        
        5-10 page detailed report with:
        - Complete requirements breakdown
        - All obligations, prohibitions, deadlines
        - Entity analysis
        - Document structure
        """
        logger.info("Generating detailed analysis DOCX...")
        
        doc = Document()
        
        # Title page
        title = doc.add_heading('Detailed Regulatory Analysis', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_heading(metadata.get('title', 'Unknown Regulation')[:100], 1)
        
        # Table of contents placeholder
        doc.add_heading('Contents', 2)
        doc.add_paragraph('1. Document Overview')
        doc.add_paragraph('2. Requirements Analysis')
        doc.add_paragraph('3. Entity Analysis')
        doc.add_paragraph('4. Document Structure')
        doc.add_paragraph('5. Compliance Summary')
        
        doc.add_page_break()
        
        # Section 1: Overview
        doc.add_heading('1. Document Overview', 1)
        
        doc.add_heading('Metadata', 2)
        for key, value in metadata.items():
            if value:
                p = doc.add_paragraph()
                p.add_run(f'{key.replace("_", " ").title()}: ').bold = True
                p.add_run(str(value)[:200])
        
        # Section 2: Requirements
        doc.add_page_break()
        doc.add_heading('2. Requirements Analysis', 1)
        
        for i, section in enumerate(requirements, 1):
            doc.add_heading(f'Section {section.get("section_number", i)}', 2)
            
            # Obligations
            if section.get('obligations'):
                doc.add_heading('Obligations', 3)
                for obligation in section['obligations']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(obligation.get('plain_language', obligation.get('text', 'N/A')))
            
            # Prohibitions
            if section.get('prohibitions'):
                doc.add_heading('Prohibitions', 3)
                for prohibition in section['prohibitions']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(prohibition.get('plain_language', prohibition.get('text', 'N/A')))
            
            # Deadlines
            if section.get('deadlines'):
                doc.add_heading('Deadlines', 3)
                for deadline in section['deadlines']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"{deadline['date']}: {deadline['action']}")
        
        # Section 3: Entity Analysis
        if entities:
            doc.add_page_break()
            doc.add_heading('3. Entity Analysis', 1)
            
            # Agencies
            if entities.get('agencies'):
                doc.add_heading('Agencies Referenced', 2)
                for agency in entities['agencies']:
                    doc.add_paragraph(
                        f"{agency['agency']}: {agency['count']} mentions",
                        style='List Bullet'
                    )
            
            # Dollar amounts
            if entities.get('dollar_amounts'):
                doc.add_heading('Key Dollar Amounts', 2)
                sorted_amounts = sorted(
                    entities['dollar_amounts'], 
                    key=lambda x: x['amount'], 
                    reverse=True
                )
                for amt in sorted_amounts[:10]:
                    doc.add_paragraph(
                        f"{amt['formatted']} - {amt['type']}",
                        style='List Bullet'
                    )
        
        # Save
        doc.save(output_path)
        logger.info(f"✅ Saved detailed analysis: {output_path}")
        
        return output_path
    
    def export_html_report(
        self,
        metadata: Dict,
        requirements: List[Dict],
        entities: Dict,
        output_path: str
    ) -> str:
        """
        Generate HTML report for web viewing.
        """
        logger.info("Generating HTML report...")
        
        # Build HTML
        html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{metadata.get('title', 'Regulatory Analysis')}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Arial, sans-serif;
            max-width: 900px;
            margin: 40px auto;
            padding: 20px;
            line-height: 1.6;
            color: #333;
        }}
        h1 {{
            color: #1a5490;
            border-bottom: 3px solid #1a5490;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #2c5f8d;
            margin-top: 30px;
            border-bottom: 1px solid #ddd;
            padding-bottom: 5px;
        }}
        h3 {{
            color: #4a7ba7;
            margin-top: 20px;
        }}
        .metadata {{
            background: #f5f7fa;
            padding: 20px;
            border-radius: 8px;
            margin: 20px 0;
        }}
        .metadata-row {{
            display: flex;
            margin: 8px 0;
        }}
        .metadata-label {{
            font-weight: bold;
            min-width: 150px;
            color: #555;
        }}
        .requirement {{
            background: #fff;
            border-left: 4px solid #1a5490;
            padding: 15px;
            margin: 15px 0;
            border-radius: 4px;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        }}
        .requirement.prohibition {{
            border-left-color: #dc3545;
        }}
        .deadline {{
            background: #fff3cd;
            border-left: 4px solid #ffc107;
            padding: 10px;
            margin: 10px 0;
            border-radius: 4px;
        }}
        .severity-high {{
            color: #dc3545;
            font-weight: bold;
        }}
        .severity-medium {{
            color: #fd7e14;
            font-weight: bold;
        }}
        .severity-low {{
            color: #28a745;
        }}
        ul {{
            margin: 10px 0;
        }}
        li {{
            margin: 8px 0;
        }}
        .stat {{
            display: inline-block;
            background: #e9ecef;
            padding: 8px 16px;
            margin: 5px;
            border-radius: 20px;
            font-size: 14px;
        }}
    </style>
</head>
<body>
    <h1>{metadata.get('title', 'Regulatory Analysis')}</h1>
    
    <div class="metadata">
        <h2>Document Information</h2>
        <div class="metadata-row">
            <span class="metadata-label">Agency:</span>
            <span>{metadata.get('agency', 'Unknown')}</span>
        </div>
        <div class="metadata-row">
            <span class="metadata-label">Type:</span>
            <span>{metadata.get('document_type', 'Unknown')}</span>
        </div>
        <div class="metadata-row">
            <span class="metadata-label">Regulation:</span>
            <span>{metadata.get('regulation_number', 'Unknown')}</span>
        </div>
        <div class="metadata-row">
            <span class="metadata-label">Publication Date:</span>
            <span>{metadata.get('publication_date', 'Unknown')}</span>
        </div>
        <div class="metadata-row">
            <span class="metadata-label">Effective Date:</span>
            <span>{metadata.get('effective_date', 'TBD')}</span>
        </div>
    </div>
    
    <div>
        <h2>Analysis Summary</h2>
"""
        
        # Add statistics
        total_obligations = sum(len(r.get('obligations', [])) for r in requirements)
        total_prohibitions = sum(len(r.get('prohibitions', [])) for r in requirements)
        total_deadlines = sum(len(r.get('deadlines', [])) for r in requirements)
        
        html += f"""
        <div>
            <span class="stat">📋 {total_obligations} Obligations</span>
            <span class="stat">🚫 {total_prohibitions} Prohibitions</span>
            <span class="stat">📅 {total_deadlines} Deadlines</span>
            <span class="stat">🏛️ {len(entities.get('agencies', []))} Agencies</span>
        </div>
    </div>
    
    <h2>Critical Requirements</h2>
"""
        
        # Add obligations
        for section in requirements:
            for obligation in section.get('obligations', []):
                severity = obligation.get('severity', 'medium')
                severity_class = f'severity-{severity}'
                
                html += f"""
    <div class="requirement">
        <div><strong>Obligation:</strong> <span class="{severity_class}">[{severity.upper()}]</span></div>
        <p>{obligation.get('plain_language', obligation.get('text', 'N/A'))}</p>
        <ul>
            <li><strong>Who must comply:</strong> {obligation.get('who_must_comply', 'Not specified')}</li>
            <li><strong>Deadline:</strong> {obligation.get('deadline', 'None specified')}</li>
        </ul>
    </div>
"""
        
        # Add prohibitions
        for section in requirements:
            for prohibition in section.get('prohibitions', []):
                html += f"""
    <div class="requirement prohibition">
        <div><strong>Prohibition:</strong></div>
        <p>{prohibition.get('plain_language', prohibition.get('text', 'N/A'))}</p>
        <ul>
            <li><strong>Prohibited action:</strong> {prohibition.get('prohibited_action', 'Not specified')}</li>
            <li><strong>Penalty:</strong> {prohibition.get('penalty', 'Not specified')}</li>
        </ul>
    </div>
"""
        
        # Add key deadlines
        html += "<h2>Key Deadlines</h2>"
        
        all_deadlines = []
        for section in requirements:
            all_deadlines.extend(section.get('deadlines', []))
        
        # Sort by date
        all_deadlines = sorted(
            [d for d in all_deadlines if d.get('date')],
            key=lambda x: x['date']
        )
        
        for deadline in all_deadlines:
            html += f"""
    <div class="deadline">
        <strong>{deadline['date']}</strong><br>
        {deadline['action']}<br>
        <small><em>Who: {deadline.get('who', 'Not specified')}</em></small>
    </div>
"""
        
        # Footer
        html += f"""
    <hr style="margin-top: 40px;">
    <p style="text-align: center; color: #666; font-size: 14px;">
        Generated by ComplianceAI • {datetime.now().strftime('%Y-%m-%d %H:%M')}
    </p>
</body>
</html>
"""
        
        # Save
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)
        
        logger.info(f"✅ Saved HTML report: {output_path}")
        return output_path
    
    def export_markdown_report(
        self,
        metadata: Dict,
        requirements: List[Dict],
        entities: Dict,
        output_path: str
    ) -> str:
        """
        Generate Markdown report.
        
        Good for documentation and GitHub.
        """
        logger.info("Generating Markdown report...")
        
        md = []
        
        # Title
        md.append(f"# {metadata.get('title', 'Regulatory Analysis')}")
        md.append("")
        
        # Metadata
        md.append("## Document Information")
        md.append("")
        md.append(f"- **Agency:** {metadata.get('agency', 'Unknown')}")
        md.append(f"- **Type:** {metadata.get('document_type', 'Unknown')}")
        md.append(f"- **Regulation:** {metadata.get('regulation_number', 'Unknown')}")
        md.append(f"- **Published:** {metadata.get('publication_date', 'Unknown')}")
        md.append(f"- **Effective:** {metadata.get('effective_date', 'TBD')}")
        md.append("")
        
        # Summary
        md.append("## Executive Summary")
        md.append("")
        md.append(metadata.get('summary', 'No summary available'))
        md.append("")
        
        # Statistics
        total_obligations = sum(len(r.get('obligations', [])) for r in requirements)
        total_prohibitions = sum(len(r.get('prohibitions', [])) for r in requirements)
        total_deadlines = sum(len(r.get('deadlines', [])) for r in requirements)
        
        md.append("## Analysis Summary")
        md.append("")
        md.append(f"- **Obligations:** {total_obligations}")
        md.append(f"- **Prohibitions:** {total_prohibitions}")
        md.append(f"- **Deadlines:** {total_deadlines}")
        md.append(f"- **Agencies Mentioned:** {len(entities.get('agencies', []))}")
        md.append("")
        
        # Requirements
        md.append("## Requirements")
        md.append("")
        
        for section in requirements:
            section_num = section.get('section_number', '?')
            md.append(f"### Section {section_num}")
            md.append("")
            
            # Obligations
            if section.get('obligations'):
                md.append("#### Obligations")
                md.append("")
                for i, obligation in enumerate(section['obligations'], 1):
                    md.append(f"{i}. **{obligation.get('plain_language', obligation.get('text', 'N/A'))}**")
                    md.append(f"   - Who: {obligation.get('who_must_comply', 'Not specified')}")
                    md.append(f"   - Severity: {obligation.get('severity', 'unknown').upper()}")
                    if obligation.get('deadline'):
                        md.append(f"   - Deadline: {obligation['deadline']}")
                    md.append("")
            
            # Deadlines
            if section.get('deadlines'):
                md.append("#### Key Dates")
                md.append("")
                for deadline in section['deadlines']:
                    md.append(f"- **{deadline['date']}:** {deadline['action']}")
                md.append("")
        
        # Save
        content = '\n'.join(md)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        logger.info(f"✅ Saved Markdown report: {output_path}")
        return output_path


def main():
    """Test export functionality."""
    
    # Sample data
    sample_metadata = {
        'agency': 'CFPB',
        'document_type': 'Final Rule',
        'regulation_number': '12 CFR Part 1002',
        'title': 'Small Business Lending Under ECOA',
        'publication_date': '2026-05-01',
        'effective_date': '2026-06-30',
        'summary': 'Revises coverage and compliance requirements for small business lending data collection.'
    }
    
    sample_requirements = [{
        'section_number': 1,
        'obligations': [{
            'text': 'Financial institutions must collect data',
            'plain_language': 'Banks must gather small business lending information',
            'who_must_comply': 'Financial institutions with 1,000+ loans',
            'severity': 'high',
            'deadline': '2028-01-01'
        }],
        'deadlines': [{
            'date': '2028-01-01',
            'action': 'Begin compliance with data collection',
            'who': 'Covered financial institutions'
        }]
    }]
    
    sample_entities = {
        'agencies': [
            {'agency': 'CFPB', 'count': 150},
            {'agency': 'FDIC', 'count': 25}
        ],
        'dollar_amounts': [
            {'amount': 1000000, 'formatted': '$1 million', 'type': 'threshold'}
        ]
    }
    
    print("\n🧪 Testing Export Utilities\n")
    print("=" * 70)
    
    exporter = ReportExporter()
    
    # Test HTML export
    print("\n1️⃣  Generating HTML report...")
    html_path = "/tmp/test_report.html"
    exporter.export_html_report(sample_metadata, sample_requirements, sample_entities, html_path)
    print(f"   ✅ Created: {html_path}")
    
    # Test Markdown export
    print("\n2️⃣  Generating Markdown report...")
    md_path = "/tmp/test_report.md"
    exporter.export_markdown_report(sample_metadata, sample_requirements, sample_entities, md_path)
    print(f"   ✅ Created: {md_path}")
    
    # Test DOCX export
    print("\n3️⃣  Generating DOCX report...")
    docx_path = "/tmp/test_report.docx"
    exporter.export_executive_summary_docx(sample_metadata, sample_requirements, sample_entities, docx_path)
    print(f"   ✅ Created: {docx_path}")
    
    print("\n" + "=" * 70)
    print("✅ All export formats working!")
    print("\nTo use with real data:")
    print("  from export_utils import ReportExporter")
    print("  exporter = ReportExporter()")
    print("  exporter.export_html_report(metadata, requirements, entities, 'output.html')")


if __name__ == "__main__":
    main()